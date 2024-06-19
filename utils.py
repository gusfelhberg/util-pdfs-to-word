import fitz  # PyMuPDF
from PIL import Image
import io
import os
from docx import Document
from docx.shared import Inches

def convert_pdf_to_images(pdf_path, dpi=300):
    pdf_document = fitz.open(pdf_path)
    images = []
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        mat = fitz.Matrix(dpi / 72, dpi / 72)  # Set the DPI
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        images.append(img)
    return images

def append_images_to_docx(images, doc):
    for img in images:
        with io.BytesIO() as output:
            img.save(output, format="PNG")
            doc.add_picture(output, width=Inches(7.3))
            # doc.add_page_break()

def set_narrow_margins(doc):
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

def append_docx_to_docx(src_docx_path, dest_doc):
    src_doc = Document(src_docx_path)
    for element in src_doc.element.body:
        dest_doc.element.body.append(element)

def process_files(file_paths, output_docx, dpi=300):
    temp_docx = "temp.docx"
    doc = Document()
    set_narrow_margins(doc)
    
    for i, file_path in enumerate(file_paths):
        print(file_path)
        if file_path.lower().endswith('.pdf'):
            images = convert_pdf_to_images(file_path, dpi)
            append_images_to_docx(images, doc)

    
        doc.save(output_docx)
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
